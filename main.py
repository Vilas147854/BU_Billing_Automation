from flask import Flask, render_template, redirect, url_for, request, flash, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_bcrypt import Bcrypt
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from datetime import datetime
import os
from calendar import monthrange, monthcalendar
from docx import Document
from models import db, User, ClassDetails, Attendance

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your_secret_key'
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
INSTANCE_DIR = os.path.join(BASE_DIR, 'instance')
os.makedirs(INSTANCE_DIR, exist_ok=True)
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{os.path.join(INSTANCE_DIR, "faculty.db")}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db.init_app(app)
bcrypt = Bcrypt(app)
login_manager = LoginManager(app)
login_manager.login_view = "login"

with app.app_context():
    db.create_all()
    if not User.query.filter_by(is_admin=True).first():
        admin = User(
            name="Admin", email="admin@example.com",
            password=bcrypt.generate_password_hash("admin123").decode('utf-8'),
            department="Admin", workload_per_week=0, is_admin=True, is_approved=True
        )
        db.session.add(admin)
        db.session.commit()

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

@app.route("/")
def home():
    return render_template("home.html")

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        name = request.form['name']
        email = request.form['email']
        password = bcrypt.generate_password_hash(request.form['password']).decode('utf-8')
        department = request.form['department']
        workload_per_week = int(request.form['workload_per_week'])

        existing_user = User.query.filter_by(email=email).first()
        if existing_user:
            flash("Email already registered. Please log in.", "danger")
            return redirect(url_for('login'))

        new_user = User(name=name, email=email, password=password, department=department, workload_per_week=workload_per_week)
        db.session.add(new_user)
        db.session.commit()
        flash("Registration submitted! Awaiting admin approval.", "success")
        return redirect(url_for('login'))

    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        user = User.query.filter_by(email=email).first()
        if user and bcrypt.check_password_hash(user.password, password):
            if not user.is_approved and not user.is_admin:
                flash("Your account is pending admin approval.", "danger")
                return redirect(url_for('login'))
            login_user(user)
            flash("Login successful!", "success")
            return redirect(url_for('dashboard'))
        else:
            flash("Invalid email or password", "danger")
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash("You have been logged out.", "info")
    return redirect(url_for('home'))

@app.route('/dashboard', methods=['GET', 'POST'])
@login_required
def dashboard():
    if request.method == 'POST':
        action = request.form.get('action')
        year = int(request.form.get('year', datetime.now().year))
        month = int(request.form.get('month', datetime.now().month))
        if action == 'generate_bill':
            return redirect(url_for('generate_bill', year=year, month=month))
        elif action == 'generate_attendance':
            return redirect(url_for('generate_attendance', year=year, month=month))
        elif action == 'filter_classes':
            classes = ClassDetails.query.filter_by(faculty_id=current_user.id).filter(
                db.extract('year', ClassDetails.date) == year,
                db.extract('month', ClassDetails.date) == month
            ).order_by(ClassDetails.date.desc()).all()
            return render_template('dashboard.html', name=current_user.name, classes=classes, selected_year=year, selected_month=month)
    
    classes = ClassDetails.query.filter_by(faculty_id=current_user.id).order_by(ClassDetails.date.desc()).all()
    return render_template('dashboard.html', name=current_user.name, classes=classes, selected_year=datetime.now().year, selected_month=datetime.now().month)

@app.route('/add_class', methods=['GET', 'POST'])
@login_required
def add_class():
    if request.method == 'POST':
        dairy_no = int(request.form['dairy_no'])
        date = datetime.strptime(request.form['date'], '%Y-%m-%d')
        subject = request.form['subject'] if request.form['subject'] != 'Other' else request.form['custom_subject']
        subject_code = request.form['subject_code'] if request.form['subject_code'] != 'Other' else request.form['custom_subject_code']
        description = request.form['description']
        actual_hours = int(request.form['actual_hours'])
        claiming_hours = int(request.form['claiming_hours'])

        # Check if dairy_no already exists
        if ClassDetails.query.filter_by(dairy_no=dairy_no).first():
            flash("Dairy number already exists. Please use a unique number.", "danger")
            return render_template('add_class.html')

        new_class = ClassDetails(
            dairy_no=dairy_no, faculty_id=current_user.id, date=date, subject=subject,
            subject_code=subject_code, description=description, 
            actual_hours=actual_hours, claiming_hours=claiming_hours
        )
        db.session.add(new_class)
        db.session.commit()
        flash("Class details added successfully!", "success")
        return redirect(url_for('dashboard'))
    return render_template('add_class.html')

@app.route('/attendance', methods=['GET', 'POST'])
@login_required
def attendance():
    if request.method == 'POST':
        if 'filter_attendance' in request.form:
            year = int(request.form.get('year', datetime.now().year))
            month = int(request.form.get('month', datetime.now().month))
            attendance_records = Attendance.query.filter_by(faculty_id=current_user.id).filter(
                db.extract('year', Attendance.date) == year,
                db.extract('month', Attendance.date) == month
            ).order_by(Attendance.date).all()
            return render_template('attendance.html', attendance=attendance_records, selected_year=year, selected_month=month)
        else:
            date = datetime.strptime(request.form['date'], '%Y-%m-%d')
            status = request.form['status']
            new_attendance = Attendance(faculty_id=current_user.id, date=date, status=status)
            db.session.add(new_attendance)
            db.session.commit()
            flash("Attendance recorded!", "success")
            return redirect(url_for('attendance'))

    attendance_records = Attendance.query.filter_by(faculty_id=current_user.id).order_by(Attendance.date).all()
    return render_template('attendance.html', attendance=attendance_records, selected_year=datetime.now().year, selected_month=datetime.now().month)

@app.route('/admin', methods=['GET', 'POST'])
@login_required
def admin():
    if not current_user.is_admin:
        flash("Access denied. Admin privileges required.", "danger")
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        if 'approve_user' in request.form:
            user_id = int(request.form['user_id'])
            user = db.session.get(User, user_id)
            user.is_approved = True
            db.session.commit()
            flash(f"User {user.name} approved!", "success")
        elif 'edit_class' in request.form:
            class_id = int(request.form['class_id'])
            cls = db.session.get(ClassDetails, class_id)
            cls.date = datetime.strptime(request.form['date'], '%Y-%m-%d')
            cls.subject = request.form['subject']
            cls.subject_code = request.form['subject_code']
            cls.description = request.form['description']
            cls.actual_hours = int(request.form['actual_hours'])
            cls.claiming_hours = int(request.form['claiming_hours'])
            db.session.commit()
            flash("Class details updated!", "success")
        elif 'edit_attendance' in request.form:
            att_id = int(request.form['att_id'])
            att = db.session.get(Attendance, att_id)
            att.date = datetime.strptime(request.form['date'], '%Y-%m-%d')
            att.status = request.form['status']
            db.session.commit()
            flash("Attendance updated!", "success")
    
    pending_users = User.query.filter_by(is_approved=False).all()
    approved_users = User.query.filter_by(is_approved=True).filter(User.is_admin == False).all()
    all_classes = ClassDetails.query.all()
    class_data = [
        {
            'id': cls.id,
            'dairy_no': cls.dairy_no,
            'faculty_name': db.session.get(User, cls.faculty_id).name,
            'date': cls.date,
            'subject': cls.subject,
            'subject_code': cls.subject_code,
            'description': cls.description,
            'actual_hours': cls.actual_hours,
            'claiming_hours': cls.claiming_hours
        } for cls in all_classes
    ]
    all_attendance = Attendance.query.all()
    attendance_data = [
        {
            'id': att.id,
            'faculty_name': db.session.get(User, att.faculty_id).name,
            'date': att.date,
            'status': att.status
        } for att in all_attendance
    ]
    
    return render_template('admin.html', pending_users=pending_users, approved_users=approved_users, classes=class_data, attendance=attendance_data)

@app.route('/generate_bill/<int:year>/<int:month>')
@login_required
def generate_bill(year, month):
    classes = ClassDetails.query.filter_by(faculty_id=current_user.id).filter(
        db.extract('year', ClassDetails.date) == year,
        db.extract('month', ClassDetails.date) == month
    ).order_by(ClassDetails.date).all()

    if not classes:
        flash("No class details found for this month.", "danger")
        return redirect(url_for('dashboard'))

    calendar = monthcalendar(year, month)
    weeks = [week for week in calendar if any(day != 0 for day in week)]

    classes_by_week = {}
    for cls in classes:
        day = cls.date.day
        for i, week in enumerate(weeks, 1):
            if day in week and day != 0:
                week_key = f"{i}{'st' if i == 1 else 'nd' if i == 2 else 'rd' if i == 3 else 'th'} Week"
                if week_key not in classes_by_week:
                    classes_by_week[week_key] = []
                classes_by_week[week_key].append(cls)
                break

    doc = Document()
    doc.add_heading('BANGALORE UNIVERSITY', level=1)
    doc.add_paragraph(f"Department: {current_user.department}")
    doc.add_paragraph(f"Workload allotted per Week: {current_user.workload_per_week} hours")
    doc.add_paragraph("ANNEXURE (Time Table)")

    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    headers = ["Sl. No", "Dairy No.", "Date", "Particulars / chapter / lectures\n(as per Time Table)\nI / II / III / IV / V / VI Sem", "Actual hours", "Claiming hours\nLab period reduced by 2/3", "Subject code"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    sl_no = 1
    for week, week_classes in sorted(classes_by_week.items()):
        row_cells = table.add_row().cells
        row_cells[0].text = week
        for i in range(1, 7):
            row_cells[i].text = ""
        for cls in week_classes:
            row_cells = table.add_row().cells
            row_cells[0].text = str(sl_no)
            row_cells[1].text = str(cls.dairy_no)
            row_cells[2].text = cls.date.strftime('%d-%m-%Y')
            row_cells[3].text = f"{cls.subject}: {cls.description}"
            row_cells[4].text = str(cls.actual_hours)
            row_cells[5].text = str(cls.claiming_hours)
            row_cells[6].text = cls.subject_code
            sl_no += 1

    filename = f"bill_{current_user.name}_{year}_{month}.docx"
    filepath = os.path.join(INSTANCE_DIR, filename)
    doc.save(filepath)
    return send_file(filepath, as_attachment=True)

@app.route('/generate_attendance/<int:year>/<int:month>')
@login_required
def generate_attendance(year, month):
    attendance = Attendance.query.filter_by(faculty_id=current_user.id).filter(
        db.extract('year', Attendance.date) == year,
        db.extract('month', Attendance.date) == month
    ).order_by(Attendance.date).all()

    if not attendance:
        flash("No attendance records found for this month.", "danger")
        return redirect(url_for('dashboard'))

    days_in_month = monthrange(year, month)[1]
    month_names = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]
    month_name = month_names[month - 1]

    doc = Document()
    doc.add_heading('DEPARTMENT OF COMPUTER SCIENCE AND APPLICATIONS', level=1)
    doc.add_paragraph(f"GUEST FACULTY ATTENDANCE EXTRACT FOR THE MONTH OF {month_name.upper()} {year}")

    table = doc.add_table(rows=2, cols=days_in_month + 1)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Name and Designation"
    for day in range(1, days_in_month + 1):
        hdr_cells[day].text = str(day)

    data_cells = table.rows[1].cells
    data_cells[0].text = f"{current_user.name}\nGuest Faculty"
    for att in attendance:
        day = att.date.day
        data_cells[day].text = att.status

    filename = f"attendance_{current_user.name}_{year}_{month}.docx"
    filepath = os.path.join(INSTANCE_DIR, filename)
    doc.save(filepath)
    return send_file(filepath, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)